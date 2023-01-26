# -*- coding: utf-8 -*-

import docx
from datetime import datetime
from PyQt5 import QtCore, QtGui, QtWidgets
from connect.connect import DataBase


class Ui_Contract(object):
    fnBtnToApplic = None
    fname = ''

    def __init__(self):
        super(Ui_Contract, self).__init__()

    def get_connect(self):
        self.db = DataBase()
        self.db.connect()

    def contract(self, widget, fnBtnToApplic):
        self.fnBtnToApplic = fnBtnToApplic
        self.fnBtnToApplic(widget)
        # self.get_connect()

        # Заголовок окна
        conractTitle = QtWidgets.QLabel(widget)
        conractTitle.setGeometry(QtCore.QRect(292, 40, 85, 30))
        conractTitle.setText("Договор")
        font = QtGui.QFont()
        font.setPointSize(15)
        font.setBold(True)
        font.setWeight(75)
        conractTitle.setFont(font)

        # Номер Договора
        numContractLayoutWidget = QtWidgets.QWidget(widget)
        numContractLayoutWidget.setGeometry(QtCore.QRect(90, 80, 480, 30))
        numContract = QtWidgets.QHBoxLayout(numContractLayoutWidget)
        numContract.setContentsMargins(0, 0, 0, 0)
        numContractLabel = QtWidgets.QLabel(numContractLayoutWidget)
        numContractLabel.setText("№ Договора: ")
        numContract.addWidget(numContractLabel)
        self.numContractLine = QtWidgets.QLineEdit(numContractLayoutWidget)
        numContract.addWidget(self.numContractLine)

        # Дата заключения Договора
        labelDate = QtWidgets.QLabel(widget)
        labelDate.setGeometry(QtCore.QRect(90, 115, 480, 30))
        labelDate.setText('Дата заключения Договора:')
        self.dateLine = QtWidgets.QDateEdit(widget)
        self.dateLine.setGeometry(QtCore.QRect(240, 120, 180, 20))
        self.dateLine.setDisplayFormat('dd.MM.yyyy')

        # Дата окончания Договора
        labelDateEnd = QtWidgets.QLabel(widget)
        labelDateEnd.setGeometry(QtCore.QRect(90, 150, 170, 20))
        labelDateEnd.setText('Дата окончания Договора:')
        self.dateLineEnd = QtWidgets.QDateEdit(widget)
        self.dateLineEnd.setGeometry(QtCore.QRect(240, 150, 180, 20))
        self.dateLineEnd.setDisplayFormat('dd.MM.yyyy')

        # Кнопка прикрепления файла
        btnFile = QtWidgets.QPushButton(widget)
        btnFile.setGeometry(QtCore.QRect(90, 180, 150, 30))
        btnFile.setText('Выбрать файл')
        btnFile.clicked.connect(lambda: self.open_file())

        # ЧекБокс
        pref = QtWidgets.QLabel(widget)
        pref.setText('или')
        pref.setGeometry(QtCore.QRect(250, 180, 30, 30))
        self.chBox = QtWidgets.QCheckBox(widget)
        self.chBox.setGeometry(QtCore.QRect(280, 180, 30, 30))
        nir = QtWidgets.QLabel(widget)
        nir.setGeometry(QtCore.QRect(295, 180, 30, 30))
        nir.setText('НИР')

        # Наименование Заказчика
        nameLayoutWidget = QtWidgets.QWidget(widget)
        nameLayoutWidget.setGeometry(QtCore.QRect(90, 210, 480, 30))
        nameClient = QtWidgets.QHBoxLayout(nameLayoutWidget)
        nameClient.setContentsMargins(0, 0, 0, 0)
        nameClientLabel = QtWidgets.QLabel(nameLayoutWidget)
        nameClientLabel.setText("Наименование Заказчика")
        nameClient.addWidget(nameClientLabel)
        self.nameClientLine = QtWidgets.QLineEdit(nameLayoutWidget)
        nameClient.addWidget(self.nameClientLine)

        # Телефон
        phoneLayoutWidget = QtWidgets.QWidget(widget)
        phoneLayoutWidget.setGeometry(QtCore.QRect(90, 240, 370, 30))
        phone = QtWidgets.QHBoxLayout(phoneLayoutWidget)
        phone.setContentsMargins(0, 0, 0, 0)
        phoneLabel = QtWidgets.QLabel(phoneLayoutWidget)
        phoneLabel.setText("Телефон: ")
        phone.addWidget(phoneLabel)
        self.phoneLine = QtWidgets.QLineEdit(phoneLayoutWidget)
        phone.addWidget(self.phoneLine)

        # Кнопка Сохранить
        btnSave = QtWidgets.QPushButton(widget)
        btnSave.setText("Сохранить")
        btnSave.setGeometry(QtCore.QRect(250, 340, 100, 30))
        btnSave.clicked.connect(lambda: self.not_save())
        # btnSave.clicked.connect(lambda: self.saveContract())

        self.textMessage = QtWidgets.QLabel(widget)
        self.textMessage.setGeometry(QtCore.QRect(180, 400, 420, 50))
        self.textMessage.setStyleSheet('color: rgb(200, 0, 0)')

        QtCore.QMetaObject.connectSlotsByName(widget)

    # Парсинг файла договора
    @QtCore.pyqtSlot()
    def open_file(self):
        try:
            fwin = QtWidgets.QMainWindow()
            self.fname = QtWidgets.QFileDialog.getOpenFileName(
                fwin, "Выбор файла", "C:/Users/Administrator/Documents",
                "Word (*.docx)")[0]
            if self.fname == '':
                text = 'Файл не выбран'
            else:
                text = f'Выбран файл {self.fname}'
                self.doc = docx.Document(self.fname)
                self.tabs = self.doc.tables
                name_client = (
                    self.tabs[0].cell(0, 0).text).split('\n')[2].strip()
                self.nameClientLine.setText(name_client)
            self.textMessage.setText(text)
            return self.fname

        except IndexError:
            self.textMessage.setText(
                'Выбранный файл не подходит \n    под формат Договора')

    def saveContract(self):
        numContract = self.numContractLine.text().strip()
        date = self.dateLine.text().strip()
        day_start = date[:2]
        month_start = date[3:5]
        year_start = date[6:]
        date = year_start + '-' + month_start + '-' + day_start
        dateEnd = self.dateLineEnd.text().strip()
        day = dateEnd[:2]
        month = dateEnd[3:5]
        year = dateEnd[6:]
        dateEnd = year + '-' + month + '-' + day
        nameClient = self.nameClientLine.text().strip()
        phone = self.phoneLine.text().strip()
        cb = self.chBox.isChecked()
        text = ''

        queryContract = '''SELECT number FROM contract'''
        numbersContract = self.db.query(queryContract)
        numbersContract = [numbersContract[i][0]
                           for i in range(len(numbersContract))]

        # проверка введённой формы
        if len(numContract) == 0:
            self.textMessage.setText("Введите номер Договора")
        elif numContract in numbersContract:
            self.textMessage.setText("Такой Договор уже существует")
        elif date == '2000-01-01':
            self.textMessage.setText("Не выбрана дата \nзаключения Договора")
        elif dateEnd == '2000-01-01':
            self.textMessage.setText("Не выбрана дата \nокончания Договора")
        elif datetime.strptime(date, "%Y-%m-%d") >= datetime.strptime(
                                                        dateEnd, "%Y-%m-%d"):
            self.textMessage.setText("Некорректная дата \nокончания Договора")
        elif self.fname == '' and cb is False:
            self.textMessage.setText(
                "Не выбран файл Договора \n   (или выберете НИР)")
        elif len(nameClient) == 0:
            self.textMessage.setText("Введите Заказчика")
        elif len(phone) == 0:
            self.textMessage.setText("Введите телефон")

        else:
            if cb is True:
                query = ('''INSERT INTO contract
                (number, client, date_start, date_end, adress, phone)
                VALUES (%s, %s, %s, %s, %s, %s)''')
                val = (numContract, nameClient, date, dateEnd,
                       '420015, РТ, г. Казань, ул. К. Маркса, д. 68',
                       phone)
                self.db.query_insert(query, val)
                last_id = self.db.query(
                    '''SELECT MAX(id) FROM contract;''')[0][0]
                query = ('''INSERT INTO applic_contract (number, date, contract_id, status)
                VALUES(%s, %s, %s, 1)''')
                val = (1, date, last_id)
                self.db.query_insert(query, val)
            else:
                query = ('''INSERT INTO contract
                (number, client, date_start, date_end, adress, phone,
                from_doc, requisites, string_for_signed)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)''')
                indx_adress = self.tabs[0].cell(
                    0, 0).text.strip().split().index('адрес:')
                index_inn = self.tabs[0].cell(
                    0, 0).text.strip().split().index('ИНН')
                adress = ' '.join(
                    self.tabs[0].cell(
                        0, 0).text.strip().split()[indx_adress+1:index_inn])
                from_doc = self.doc.paragraphs[7].text
                requisites = self.tabs[0].cell(0, 0).text
                string_for_signed = self.tabs[2].cell(1, 0).text
                val = (numContract, nameClient, date, dateEnd, adress, phone,
                       from_doc, requisites, string_for_signed)
                self.db.query_insert(query, val)
                queryLog = f'Внесён Договор №{numContract}'
                self.db.insert_log(queryLog)

                try:
                    query = ('''INSERT INTO applic_contract (
                        number, date, contract_id, samples_rp, target,
                        status, price, nds)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)''')
                    last_id = self.db.query(
                        '''SELECT MAX(id) FROM contract;''')[0][0]
                    samples_rp = self.doc.paragraphs[170].text.split(
                        'провести исследования')[1].strip('.')
                    target = self.doc.paragraphs[173].text.split(
                        'является определение')[1].replace('.', '')
                    a = len(self.tabs[1].column_cells(0))
                    price, nds = self.tabs[1].cell(a-1, 0).text.split('\n')
                    price = float(price.split(
                        'Итого  общая сумма заявки:')[1].strip().replace(
                        ' ', '').replace(',', '.').replace("\xa0", ''))
                    nds = float(nds.split(
                        'С учетом НДС 20%:')[1].strip().replace(
                        ' ', '').replace(',', '.').replace("\xa0", ''))
                    val = (1, date, last_id, samples_rp, target, 1, price, nds)
                    self.db.query_insert(query, val)
                    last_id_ac = self.db.query(
                        '''SELECT MAX(id) FROM applic_contract;''')[0][0]
                    query = ('''INSERT INTO types_of_research(type_research,
                        count_sample, deadline, price, applic_contract_id)
                        VALUES (%s, %s, %s, %s, %s)''')

                    for i in range(1, a-1):
                        if i != a:
                            val = (self.tabs[1].cell(i, 0).text,
                                   self.tabs[1].cell(i, 1).text,
                                   self.tabs[1].cell(i, 2).text,
                                   float(self.tabs[1].cell(i, 3).text.replace(
                                    ' ', ''
                                    ).replace(',', '.').replace("\xa0", '')),
                                   last_id_ac)
                            self.db.query_insert(query, val)
                    queryLog = f'Внесена Заявка №1 к Договору №{numContract}'
                    self.db.insert_log(queryLog)

                except ValueError:
                    text = f'''\nЗаявка №1 по Договору{
                        numContract
                    } не сохранена \nВнесете вручную'''
                except IndexError:
                    text = f'''\nЗаявка №1 по Договору{
                        numContract
                    } не сохранена \nВнесете вручную'''

            self.textMessage.setText(f"Сохранён Договор №{numContract} {text}")
            self.clearFilds()

    def clearFilds(self):
        '''Очистка полей'''
        self.numContractLine.setText('')
        self.dateLine.setDate(QtCore.QDate(2000, 1, 1))
        self.dateLineEnd.setDate(QtCore.QDate(2000, 1, 1))
        self.fname = ''
        self.chBox.setChecked(False)
        self.nameClientLine.setText('')
        self.phoneLine.setText('')

    def not_save(self):
        self.textMessage.setText('БД не подключена')
